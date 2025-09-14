#!/usr/bin/env python3
"""
File Preprocessing Script for College Buddy RAG System

This script converts various file formats (.doc, .docx, .html) to clean .txt files
in the backend/data directory for ingestion into the vector database.

Requirements:
- python-docx for .docx files
- win32com.client for .doc files (Windows only)
- beautifulsoup4 for .html files
"""

import os
import re
from pathlib import Path
from typing import List, Optional

# Import required libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. .docx files will be skipped.")

try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    print("Warning: pywin32 not installed. .doc files will be skipped.")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    print("Warning: beautifulsoup4 not installed. .html files will be skipped.")


def clean_text(text: str) -> str:
    """
    Clean text by removing BOM, extra whitespace, and artifacts.
    
    Args:
        text: Raw text content
        
    Returns:
        Cleaned text content
    """
    if not text:
        return ""
    
    # Remove BOM and other Unicode artifacts
    text = text.replace('\ufeff', '')  # BOM
    text = text.replace('ï»¿', '')     # UTF-8 BOM
    text = text.replace('\u200b', '')  # Zero-width space
    
    # Remove extra whitespace and normalize line breaks
    text = re.sub(r'\r\n', '\n', text)  # Normalize Windows line endings
    text = re.sub(r'\r', '\n', text)    # Normalize Mac line endings
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Remove excessive line breaks
    text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces and tabs
    text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)  # Trim lines
    
    # Remove common HTML artifacts that might slip through
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&quot;', '"', text)
    
    return text.strip()


def convert_doc_to_docx(doc_path: Path) -> Optional[Path]:
    """
    Convert .doc file to .docx using win32com.client.
    
    Args:
        doc_path: Path to the .doc file
        
    Returns:
        Path to the converted .docx file, or None if conversion failed
    """
    if not WIN32_AVAILABLE:
        print(f"Skipping {doc_path.name}: win32com.client not available")
        return None
    
    try:
        # Create Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open the .doc file
        doc = word.Documents.Open(str(doc_path.absolute()))
        
        # Create output path
        docx_path = doc_path.with_suffix('.docx')
        
        # Save as .docx
        doc.SaveAs2(str(docx_path.absolute()), FileFormat=16)  # 16 = wdFormatXMLDocument
        doc.Close()
        word.Quit()
        
        print(f"Converted: {doc_path.name} → {docx_path.name}")
        return docx_path
        
    except Exception as e:
        print(f"Error converting {doc_path.name}: {e}")
        return None


def extract_text_from_docx(docx_path: Path) -> str:
    """
    Extract text from .docx file using python-docx.
    
    Args:
        docx_path: Path to the .docx file
        
    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        return ""
    
    try:
        doc = Document(docx_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        print(f"Error reading {docx_path.name}: {e}")
        return ""


def extract_text_from_html(html_path: Path) -> str:
    """
    Extract text from .html file using BeautifulSoup.
    
    Args:
        html_path: Path to the .html file
        
    Returns:
        Extracted text content
    """
    if not BS4_AVAILABLE:
        return ""
    
    try:
        with open(html_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        return text
        
    except Exception as e:
        print(f"Error reading {html_path.name}: {e}")
        return ""


def process_file(file_path: Path, data_dir: Path) -> bool:
    """
    Process a single file and convert it to .txt format.
    
    Args:
        file_path: Path to the file to process
        data_dir: Directory containing the data files
        
    Returns:
        True if processing was successful, False otherwise
    """
    print(f"Processing: {file_path.name}")
    
    text_content = ""
    success = False
    
    if file_path.suffix.lower() == '.doc':
        # Convert .doc to .docx first
        docx_path = convert_doc_to_docx(file_path)
        if docx_path and docx_path.exists():
            text_content = extract_text_from_docx(docx_path)
            # Clean up temporary .docx file
            try:
                docx_path.unlink()
                print(f"Cleaned up temporary file: {docx_path.name}")
            except:
                pass
            success = bool(text_content.strip())
        else:
            print(f"Failed to convert {file_path.name}")
            
    elif file_path.suffix.lower() == '.docx':
        text_content = extract_text_from_docx(file_path)
        success = bool(text_content.strip())
        
    elif file_path.suffix.lower() == '.html':
        text_content = extract_text_from_html(file_path)
        success = bool(text_content.strip())
    
    if success and text_content.strip():
        # Clean the text
        cleaned_text = clean_text(text_content)
        
        if cleaned_text.strip():
            # Create output .txt file
            txt_path = data_dir / f"{file_path.stem}.txt"
            
            try:
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(cleaned_text)
                print(f"Created: {txt_path.name}")
                return True
            except Exception as e:
                print(f"Error writing {txt_path.name}: {e}")
        else:
            print(f"No content extracted from {file_path.name}")
    else:
        print(f"Failed to extract content from {file_path.name}")
    
    return False


def main():
    """Main function to process all files in the data directory."""
    # Get the data directory
    script_dir = Path(__file__).parent
    data_dir = script_dir / "data"
    
    if not data_dir.exists():
        print(f"Data directory not found: {data_dir}")
        print("Please create the backend/data directory and add your files.")
        return
    
    # Find all files to process
    supported_extensions = ['.doc', '.docx', '.html']
    files_to_process = []
    
    for ext in supported_extensions:
        files_to_process.extend(data_dir.glob(f"*{ext}"))
        files_to_process.extend(data_dir.glob(f"*{ext.upper()}"))
    
    if not files_to_process:
        print("No .doc, .docx, or .html files found in the data directory.")
        return
    
    print(f"Found {len(files_to_process)} files to process:")
    for file_path in files_to_process:
        print(f"  - {file_path.name}")
    print()
    
    # Process each file
    successful_conversions = 0
    
    for file_path in files_to_process:
        if process_file(file_path, data_dir):
            successful_conversions += 1
        print()
    
    print(f"Processing complete!")
    print(f"Successfully converted: {successful_conversions}/{len(files_to_process)} files")
    
    if successful_conversions > 0:
        print("\nYou can now run 'python ingest.py' to build the vector database.")
    else:
        print("\nNo files were successfully converted. Please check the error messages above.")


if __name__ == "__main__":
    main()
